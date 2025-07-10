"""
AI Knowledge Base Service
Handles document upload, processing, embedding, and retrieval for AI calls
"""

import os
import json
import logging
import hashlib
from datetime import datetime
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
import uuid

from config import get_config
from database import get_supabase_client

logger = logging.getLogger(__name__)
config = get_config()
supabase = get_supabase_client()

@dataclass
class Document:
    id: str
    title: str
    content: str
    file_type: str
    file_size: int
    upload_date: datetime
    user_id: str
    tags: List[str]
    summary: str
    chunk_count: int

@dataclass
class DocumentChunk:
    id: str
    document_id: str
    content: str
    chunk_index: int
    embedding: List[float]
    token_count: int

class DocumentProcessor:
    """Document processing and text extraction"""
    
    def __init__(self):
        self.supported_types = [
            'text/plain', 'text/markdown', 'text/csv',
            'application/json', 'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        ]
    
    def extract_text(self, file_content: bytes, file_type: str, filename: str) -> str:
        """Extract text content from various file types"""
        try:
            if file_type in ['text/plain', 'text/markdown', 'text/csv']:
                return file_content.decode('utf-8')
            
            elif file_type == 'application/json':
                data = json.loads(file_content.decode('utf-8'))
                return self._json_to_text(data)
            
            elif file_type == 'application/pdf':
                return self._extract_pdf_text(file_content)
            
            elif 'word' in file_type or 'document' in file_type:
                return self._extract_docx_text(file_content)
            
            elif 'spreadsheet' in file_type or 'excel' in file_type:
                return self._extract_xlsx_text(file_content)
            
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            raise
    
    def _json_to_text(self, data: Any, prefix: str = "") -> str:
        """Convert JSON data to readable text"""
        text_parts = []
        
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    text_parts.append(f"{prefix}{key}:")
                    text_parts.append(self._json_to_text(value, prefix + "  "))
                else:
                    text_parts.append(f"{prefix}{key}: {str(value)}")
        
        elif isinstance(data, list):
            for i, item in enumerate(data):
                text_parts.append(f"{prefix}Item {i+1}:")
                text_parts.append(self._json_to_text(item, prefix + "  "))
        
        else:
            text_parts.append(f"{prefix}{str(data)}")
        
        return "\n".join(text_parts)
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF files"""
        try:
            import PyPDF2
            from io import BytesIO
            
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text_parts = []
            
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
            
            return "\n".join(text_parts)
            
        except ImportError:
            # Fallback if PyPDF2 not available
            logger.warning("PyPDF2 not available, using basic text extraction")
            try:
                return file_content.decode('utf-8', errors='ignore')
            except:
                return "PDF content could not be extracted. Please install PyPDF2 for better PDF support."
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX files"""
        try:
            import docx
            from io import BytesIO
            
            doc = docx.Document(BytesIO(file_content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            return "\n".join(text_parts)
            
        except ImportError:
            logger.warning("python-docx not available, using basic text extraction")
            try:
                return file_content.decode('utf-8', errors='ignore')
            except:
                return "DOCX content could not be extracted. Please install python-docx for better document support."
    
    def _extract_xlsx_text(self, file_content: bytes) -> str:
        """Extract text from XLSX files"""
        try:
            import openpyxl
            from io import BytesIO
            
            workbook = openpyxl.load_workbook(BytesIO(file_content))
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
            
        except ImportError:
            logger.warning("openpyxl not available, using basic text extraction")
            try:
                return file_content.decode('utf-8', errors='ignore')
            except:
                return "XLSX content could not be extracted. Please install openpyxl for better spreadsheet support."

class TextChunker:
    """Split text into chunks for embedding"""
    
    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def chunk_text(self, text: str, document_id: str) -> List[DocumentChunk]:
        """Split text into overlapping chunks"""
        chunks = []
        words = text.split()
        
        for i in range(0, len(words), self.chunk_size - self.overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = " ".join(chunk_words)
            
            if chunk_text.strip():
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=document_id,
                    content=chunk_text,
                    chunk_index=len(chunks),
                    embedding=[],  # Will be populated by embedding service
                    token_count=len(chunk_words)
                )
                chunks.append(chunk)
        
        return chunks

class EmbeddingService:
    """Generate embeddings for text chunks"""
    
    def __init__(self):
        self.model_name = "text-embedding-ada-002"  # OpenAI embedding model
    
    def generate_embedding(self, text: str) -> List[float]:
        """Generate embedding for text using OpenAI API"""
        try:
            # Note: This requires openai library and API key
            import openai
            
            # Get API key from environment
            openai_api_key = os.getenv('OPENAI_API_KEY')
            if not openai_api_key:
                raise ValueError("OPENAI_API_KEY environment variable not set")
            
            client = openai.OpenAI(api_key=openai_api_key)
            
            response = client.embeddings.create(
                model=self.model_name,
                input=text
            )
            
            return response.data[0].embedding
            
        except ImportError:
            logger.warning("OpenAI library not available, using mock embeddings")
            return self._generate_mock_embedding(text)
        except Exception as e:
            logger.error(f"Error generating embedding: {e}")
            return self._generate_mock_embedding(text)
    
    def _generate_mock_embedding(self, text: str) -> List[float]:
        """Generate mock embedding for development/testing"""
        # Create a simple hash-based mock embedding
        text_hash = hashlib.md5(text.encode()).hexdigest()
        
        # Convert hex to float values (1536 dimensions like OpenAI)
        embedding = []
        for i in range(0, len(text_hash), 2):
            hex_pair = text_hash[i:i+2]
            float_val = int(hex_pair, 16) / 255.0 - 0.5  # Normalize to [-0.5, 0.5]
            embedding.append(float_val)
        
        # Pad to 1536 dimensions
        while len(embedding) < 1536:
            embedding.extend(embedding[:min(32, 1536 - len(embedding))])
        
        return embedding[:1536]

class KnowledgeBase:
    """Main knowledge base service"""
    
    def __init__(self, user_id: str):
        self.user_id = user_id
        self.processor = DocumentProcessor()
        self.chunker = TextChunker()
        self.embedding_service = EmbeddingService()
    
    def upload_document(self, file_content: bytes, filename: str, file_type: str, tags: List[str] = None) -> Dict[str, Any]:
        """Upload and process a document"""
        try:
            # Extract text content
            text_content = self.processor.extract_text(file_content, file_type, filename)
            
            if not text_content.strip():
                raise ValueError("Document contains no extractable text")
            
            # Generate document summary
            summary = self._generate_summary(text_content)
            
            # Create document record
            document_id = str(uuid.uuid4())
            document = Document(
                id=document_id,
                title=filename,
                content=text_content,
                file_type=file_type,
                file_size=len(file_content),
                upload_date=datetime.utcnow(),
                user_id=self.user_id,
                tags=tags or [],
                summary=summary,
                chunk_count=0
            )
            
            # Save document to database
            doc_data = {
                "id": document.id,
                "title": document.title,
                "content": document.content,
                "file_type": document.file_type,
                "file_size": document.file_size,
                "upload_date": document.upload_date.isoformat(),
                "user_id": document.user_id,
                "tags": json.dumps(document.tags),
                "summary": document.summary,
                "chunk_count": 0,
                "created_at": datetime.utcnow().isoformat(),
                "updated_at": datetime.utcnow().isoformat()
            }
            
            supabase.table("knowledge_documents").insert(doc_data).execute()
            
            # Process chunks
            chunks = self.chunker.chunk_text(text_content, document_id)
            processed_chunks = []
            
            for chunk in chunks:
                # Generate embedding
                chunk.embedding = self.embedding_service.generate_embedding(chunk.content)
                
                # Save chunk to database
                chunk_data = {
                    "id": chunk.id,
                    "document_id": chunk.document_id,
                    "content": chunk.content,
                    "chunk_index": chunk.chunk_index,
                    "embedding": json.dumps(chunk.embedding),
                    "token_count": chunk.token_count,
                    "created_at": datetime.utcnow().isoformat()
                }
                
                supabase.table("knowledge_chunks").insert(chunk_data).execute()
                processed_chunks.append(chunk)
            
            # Update document with chunk count
            supabase.table("knowledge_documents").update({
                "chunk_count": len(processed_chunks)
            }).eq("id", document_id).execute()
            
            logger.info(f"Successfully processed document {filename} with {len(processed_chunks)} chunks")
            
            return {
                "status": "success",
                "document_id": document_id,
                "title": filename,
                "chunk_count": len(processed_chunks),
                "summary": summary,
                "file_size": len(file_content),
                "message": f"Document processed successfully with {len(processed_chunks)} chunks"
            }
            
        except Exception as e:
            logger.exception(f"Error uploading document {filename}")
            return {
                "status": "error",
                "message": str(e)
            }
    
    def search_knowledge(self, query: str, max_results: int = 5) -> List[Dict[str, Any]]:
        """Search knowledge base using semantic similarity"""
        try:
            # Generate embedding for query
            query_embedding = self.embedding_service.generate_embedding(query)
            
            # Get all chunks for user
            chunks_result = supabase.table("knowledge_chunks").select("*").eq("user_id", self.user_id).execute()
            chunks = chunks_result.data or []
            
            if not chunks:
                return []
            
            # Calculate similarity scores
            scored_chunks = []
            for chunk in chunks:
                try:
                    chunk_embedding = json.loads(chunk["embedding"])
                    similarity = self._cosine_similarity(query_embedding, chunk_embedding)
                    
                    scored_chunks.append({
                        "chunk_id": chunk["id"],
                        "document_id": chunk["document_id"],
                        "content": chunk["content"],
                        "similarity": similarity,
                        "chunk_index": chunk["chunk_index"]
                    })
                except Exception as e:
                    logger.warning(f"Error processing chunk {chunk['id']}: {e}")
                    continue
            
            # Sort by similarity and return top results
            scored_chunks.sort(key=lambda x: x["similarity"], reverse=True)
            
            # Get document titles for results
            results = []
            for chunk in scored_chunks[:max_results]:
                doc_result = supabase.table("knowledge_documents").select("title, summary").eq("id", chunk["document_id"]).maybe_single().execute()
                
                document_title = doc_result.data.get("title", "Unknown") if doc_result.data else "Unknown"
                document_summary = doc_result.data.get("summary", "") if doc_result.data else ""
                
                results.append({
                    "content": chunk["content"],
                    "document_title": document_title,
                    "document_summary": document_summary,
                    "similarity": chunk["similarity"],
                    "source": f"{document_title} (chunk {chunk['chunk_index'] + 1})"
                })
            
            return results
            
        except Exception as e:
            logger.exception("Error searching knowledge base")
            return []
    
    def get_user_documents(self) -> List[Dict[str, Any]]:
        """Get all documents for user"""
        try:
            result = supabase.table("knowledge_documents").select("*").eq("user_id", self.user_id).order("created_at", desc=True).execute()
            
            documents = []
            for doc in result.data or []:
                documents.append({
                    "id": doc["id"],
                    "title": doc["title"],
                    "file_type": doc["file_type"],
                    "file_size": doc["file_size"],
                    "upload_date": doc["upload_date"],
                    "tags": json.loads(doc.get("tags", "[]")),
                    "summary": doc["summary"],
                    "chunk_count": doc["chunk_count"],
                    "created_at": doc["created_at"]
                })
            
            return documents
            
        except Exception as e:
            logger.exception("Error getting user documents")
            return []
    
    def delete_document(self, document_id: str) -> bool:
        """Delete a document and its chunks"""
        try:
            # Delete chunks first
            supabase.table("knowledge_chunks").delete().eq("document_id", document_id).execute()
            
            # Delete document
            result = supabase.table("knowledge_documents").delete().eq("id", document_id).eq("user_id", self.user_id).execute()
            
            return bool(result.data)
            
        except Exception as e:
            logger.exception(f"Error deleting document {document_id}")
            return False
    
    def _generate_summary(self, text: str, max_length: int = 200) -> str:
        """Generate a summary of the document"""
        # Simple extractive summarization
        sentences = text.replace('\n', ' ').split('. ')
        
        if len(sentences) <= 3:
            return text[:max_length]
        
        # Take first few sentences up to max_length
        summary = ""
        for sentence in sentences[:5]:
            if len(summary + sentence) < max_length:
                summary += sentence + ". "
            else:
                break
        
        return summary.strip()
    
    def _cosine_similarity(self, vec1: List[float], vec2: List[float]) -> float:
        """Calculate cosine similarity between two vectors"""
        try:
            import math
            
            # Ensure vectors are same length
            min_len = min(len(vec1), len(vec2))
            vec1 = vec1[:min_len]
            vec2 = vec2[:min_len]
            
            # Calculate dot product
            dot_product = sum(a * b for a, b in zip(vec1, vec2))
            
            # Calculate magnitudes
            magnitude1 = math.sqrt(sum(a * a for a in vec1))
            magnitude2 = math.sqrt(sum(a * a for a in vec2))
            
            if magnitude1 == 0 or magnitude2 == 0:
                return 0
            
            return dot_product / (magnitude1 * magnitude2)
            
        except Exception as e:
            logger.warning(f"Error calculating similarity: {e}")
            return 0

# AI Function for Vapi Integration
def ai_search_knowledge(user_id: str, query: str, max_results: str = "3") -> Dict[str, Any]:
    """AI Function: Search knowledge base during calls"""
    try:
        kb = KnowledgeBase(user_id)
        results = kb.search_knowledge(query, int(max_results))
        
        if results:
            # Format results for AI response
            formatted_results = []
            for result in results:
                formatted_results.append({
                    "content": result["content"][:500] + "..." if len(result["content"]) > 500 else result["content"],
                    "source": result["source"],
                    "relevance": f"{result['similarity']:.2f}"
                })
            
            return {
                "status": "success",
                "results": formatted_results,
                "total_found": len(results),
                "message": f"Found {len(results)} relevant documents"
            }
        else:
            return {
                "status": "success",
                "results": [],
                "total_found": 0,
                "message": "No relevant documents found in knowledge base"
            }
            
    except Exception as e:
        logger.exception("Error in AI search knowledge")
        return {
            "status": "error",
            "message": f"Error searching knowledge base: {str(e)}"
        }